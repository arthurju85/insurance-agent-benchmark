#!/usr/bin/env python3
"""
生成 Word 版本的需求文档
需要安装: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = parse_xml(r'<w:{} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'.format(edge))
                    tcBorders.append(element)
                for key in ["sz", "val", "color", "space"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_prd_word():
    """创建 PRD Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)

    # 标题
    title = doc.add_heading('保险智能体评测系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('产品需求文档 (PRD)', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 文档信息
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ('文档版本', 'v1.0.0'),
        ('创建日期', '2026-02-24'),
        ('最后更新', '2026-02-24'),
        ('文档状态', '已完成开发')
    ]

    for i, (key, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # 目录
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 文档信息',
        '2. 产品概述',
        '3. 功能需求',
        '4. 非功能需求',
        '5. 系统架构',
        '6. API 规范',
        '7. 数据模型',
        '8. 测试需求',
        '9. 部署需求',
        '10. 待办事项',
        '11. 附录'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # 1. 文档信息
    doc.add_heading('1. 文档信息', level=1)

    doc.add_heading('1.1 版本历史', level=2)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    headers = ['版本', '日期', '修改人', '修改内容', '状态']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    row = table.rows[1]
    row.cells[0].text = 'v1.0.0'
    row.cells[1].text = '2026-02-24'
    row.cells[2].text = 'Claude'
    row.cells[3].text = '初始版本，完成所有P1/P2功能'
    row.cells[4].text = '已完成'

    doc.add_heading('1.2 相关文档', level=2)
    doc.add_paragraph('• 技术设计文档: docs/TDD.md', style='List Bullet')
    doc.add_paragraph('• API 接口文档: http://localhost:8000/docs', style='List Bullet')
    doc.add_paragraph('• 部署文档: docs/DEPLOY.md', style='List Bullet')
    doc.add_paragraph('• 用户手册: docs/USER_GUIDE.md', style='List Bullet')

    # 2. 产品概述
    doc.add_heading('2. 产品概述', level=1)

    doc.add_heading('2.1 产品背景', level=2)
    doc.add_paragraph(
        '随着大语言模型在保险行业的应用，需要一个专业的评测平台来评估不同保险Agent的能力表现，'
        '确保其具备足够的专业知识、推理能力和合规意识。'
    )

    doc.add_heading('2.2 产品目标', level=2)
    goals = [
        '标准化评测: 建立统一的保险Agent能力评测标准',
        '多维度评估: 从知识、理解、推理、合规、工具五个维度全面评估',
        '数据驱动: 通过排行榜和历史数据分析Agent能力演进',
        '防污染设计: 题目变异引擎防止Agent过拟合'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('2.3 目标用户', level=2)
    user_table = doc.add_table(rows=5, cols=3)
    user_table.style = 'Table Grid'
    headers = ['用户角色', '使用场景', '核心需求']
    for i, header in enumerate(headers):
        user_table.rows[0].cells[i].text = header
        user_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    users = [
        ('保险公司', '评估内部Agent产品', '了解产品能力边界、发现改进点'),
        ('技术厂商', '展示产品能力', '获得权威评测认证、生成排行榜'),
        ('监管机构', '监督AI应用合规性', '确保Agent符合行业规范'),
        ('研究人员', '学术研究', '获取评测数据集、对比分析')
    ]
    for i, (role, scenario, need) in enumerate(users, 1):
        user_table.rows[i].cells[0].text = role
        user_table.rows[i].cells[1].text = scenario
        user_table.rows[i].cells[2].text = need

    # 3. 功能需求
    doc.add_page_break()
    doc.add_heading('3. 功能需求', level=1)

    doc.add_heading('3.1 P1 优先级功能（核心功能）', level=2)

    # 3.1.1 核心评测流水线
    doc.add_heading('3.1.1 核心评测流水线 [已完成]', level=3)
    doc.add_paragraph('需求描述: 支持五维度自动化评测，支持多种题型。')
    doc.add_paragraph('功能点:', style='List Bullet')
    doc.add_paragraph('五维度评测（知识/理解/推理/合规/工具）', style='List Bullet 2')
    doc.add_paragraph('多线程并行执行', style='List Bullet 2')
    doc.add_paragraph('自动评分和报告生成', style='List Bullet 2')
    doc.add_paragraph('支持多种题型（选择/填空/计算/案例分析/多轮对话/工具调用）', style='List Bullet 2')

    doc.add_paragraph('验收标准:', style='List Bullet')
    doc.add_paragraph('单次评测可并发执行 ≥5 道题', style='List Bullet 2')
    doc.add_paragraph('评测结果自动保存到数据库', style='List Bullet 2')
    doc.add_paragraph('支持 JSON 格式报告导出', style='List Bullet 2')

    doc.add_paragraph('相关代码: backend/src/pipeline/')

    # 3.1.2 沙箱系统
    doc.add_heading('3.1.2 沙箱系统 [已完成]', level=3)
    doc.add_paragraph('需求描述: 提供隔离的Agent执行环境，支持主流大模型API。')
    doc.add_paragraph('功能点:', style='List Bullet')
    doc.add_paragraph('隔离执行环境（超时/异常处理）', style='List Bullet 2')
    doc.add_paragraph('OpenAI API 适配器', style='List Bullet 2')
    doc.add_paragraph('Azure OpenAI 适配器', style='List Bullet 2')
    doc.add_paragraph('本地模型适配器（vLLM）', style='List Bullet 2')
    doc.add_paragraph('工具调用支持', style='List Bullet 2')

    # 3.1.3 数据持久化
    doc.add_heading('3.1.3 数据持久化 [已完成]', level=3)
    doc.add_paragraph('需求描述: 评测结果持久化存储，支持历史查询和趋势分析。')

    # 数据表结构
    doc.add_paragraph('数据表结构:', style='List Bullet')
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['表名', '用途', '关键字段']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    tables = [
        ('evaluations', '评测结果主表', 'id, agent_id, score, status'),
        ('evaluation_details', '详细评分', 'question_id, score, latency'),
        ('leaderboard_history', '排行榜历史', 'rank, agent_id, date'),
        ('arena_sessions', '竞技场会话', 'session_id, status'),
        ('arena_events', '事件日志', 'event_type, timestamp'),
        ('registered_agents', 'Agent注册', 'agent_id, config, is_active')
    ]
    for i, (name, purpose, fields) in enumerate(tables, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = purpose
        table.rows[i].cells[2].text = fields

    # 3.1.4 题目变异引擎
    doc.add_heading('3.1.4 题目变异引擎 [已完成]', level=3)
    doc.add_paragraph('需求描述: 自动生成题目变体，防止数据污染和Agent过拟合。')

    doc.add_paragraph('变异策略:', style='List Bullet')
    variant_table = doc.add_table(rows=5, cols=3)
    variant_table.style = 'Table Grid'
    headers = ['策略', '示例', '约束条件']
    for i, header in enumerate(headers):
        variant_table.rows[0].cells[i].text = header
        variant_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    variants = [
        ('日期变异', '2024-01-01 → 2023-12-17', '保持时间间隔不变'),
        ('金额变异', '10,000元 → 9,710元', '变异幅度 ±20%'),
        ('实体替换', '张某 → 李某', '同一次变异保持一致'),
        ('句式重组', '"因为...所以..." → "由于...因此..."', '语义不变')
    ]
    for i, (strategy, example, constraint) in enumerate(variants, 1):
        variant_table.rows[i].cells[0].text = strategy
        variant_table.rows[i].cells[1].text = example
        variant_table.rows[i].cells[2].text = constraint

    doc.add_page_break()

    # 3.2 P2 功能
    doc.add_heading('3.2 P2 优先级功能（增强功能）', level=2)

    doc.add_heading('3.2.1 数据爬虫系统 [已完成]', level=3)
    doc.add_paragraph('功能点:', style='List Bullet')
    doc.add_paragraph('银保监会政策爬虫', style='List Bullet 2')
    doc.add_paragraph('保险公司条款爬虫', style='List Bullet 2')
    doc.add_paragraph('裁判文书网案例爬虫', style='List Bullet 2')
    doc.add_paragraph('条款结构化解析', style='List Bullet 2')

    doc.add_heading('3.2.2 管理员后台 [已完成]', level=3)
    admin_table = doc.add_table(rows=5, cols=3)
    admin_table.style = 'Table Grid'
    headers = ['模块', '功能', '状态']
    for i, header in enumerate(headers):
        admin_table.rows[0].cells[i].text = header
        admin_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    admin_features = [
        ('仪表盘', '数据统计、趋势分析', '✅ 已完成'),
        ('题目管理', 'CRUD、变体生成、批量导入', '✅ 已完成'),
        ('评测监控', '实时监控、历史查询、性能分析', '✅ 已完成'),
        ('系统管理', '配置管理、数据备份', '✅ 已完成')
    ]
    for i, (module, features, status) in enumerate(admin_features, 1):
        admin_table.rows[i].cells[0].text = module
        admin_table.rows[i].cells[1].text = features
        admin_table.rows[i].cells[2].text = status

    # 4. 非功能需求
    doc.add_page_break()
    doc.add_heading('4. 非功能需求', level=1)

    doc.add_heading('4.1 性能需求', level=2)
    perf_table = doc.add_table(rows=5, cols=3)
    perf_table.style = 'Table Grid'
    headers = ['指标', '目标值', '测试方法']
    for i, header in enumerate(headers):
        perf_table.rows[0].cells[i].text = header
        perf_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    perfs = [
        ('API响应时间', '< 200ms (P95)', '压测工具'),
        ('评测并发数', '≥ 5 道题目并行', '实际测试'),
        ('数据库查询', '< 100ms', 'SQL分析'),
        ('前端加载', '< 3s (首屏)', 'Lighthouse')
    ]
    for i, (metric, target, method) in enumerate(perfs, 1):
        perf_table.rows[i].cells[0].text = metric
        perf_table.rows[i].cells[1].text = target
        perf_table.rows[i].cells[2].text = method

    doc.add_heading('4.2 安全需求', level=2)
    security_items = [
        '✅ API 请求支持 CORS 配置',
        '⬜ API 认证机制 (JWT/OAuth2) - 待实现',
        '⬜ Agent API Key 加密存储 - 待实现',
        '⬜ SQL 注入防护 - 待实现',
        '⬜ XSS 防护 - 待实现'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    # 5. 系统架构
    doc.add_page_break()
    doc.add_heading('5. 系统架构', level=1)

    doc.add_heading('5.1 技术栈', level=2)
    tech_table = doc.add_table(rows=5, cols=3)
    tech_table.style = 'Table Grid'
    headers = ['层级', '技术', '版本']
    for i, header in enumerate(headers):
        tech_table.rows[0].cells[i].text = header
        tech_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    techs = [
        ('前端', 'Next.js + React', '^14'),
        ('后端', 'FastAPI', '^0.109'),
        ('数据库', 'SQLite', '-'),
        ('Python', 'Python', '3.10+')
    ]
    for i, (layer, tech, version) in enumerate(techs, 1):
        tech_table.rows[i].cells[0].text = layer
        tech_table.rows[i].cells[1].text = tech
        tech_table.rows[i].cells[2].text = version

    # 6. API 规范
    doc.add_page_break()
    doc.add_heading('6. API 规范', level=1)

    doc.add_heading('6.1 基础信息', level=2)
    doc.add_paragraph('Base URL: http://localhost:8000/api/v1', style='List Bullet')
    doc.add_paragraph('数据格式: JSON', style='List Bullet')
    doc.add_paragraph('字符编码: UTF-8', style='List Bullet')
    doc.add_paragraph('认证方式: 暂无（待实现 JWT）', style='List Bullet')

    doc.add_heading('6.2 响应格式', level=2)
    doc.add_paragraph('''
{
  "success": true,
  "data": {},
  "message": "操作成功",
  "timestamp": "2026-02-24T10:00:00Z"
}
    ''')

    # 7. 数据模型
    doc.add_page_break()
    doc.add_heading('7. 数据模型', level=1)

    doc.add_heading('7.1 Question（题目）', level=2)
    doc.add_paragraph('''
{
  "question_id": "string",      // 题目唯一ID
  "dimension": "string",        // 评测维度
  "question_type": "string",    // 题目类型
  "difficulty": "string",       // 难度等级
  "title": "string",            // 题目标题
  "content": "string",          // 题目内容
  "context": "string",          // 上下文信息
  "ground_truth": {},           // 标准答案
  "validation_rules": {},       // 验证规则
  "score": 100.0,               // 分值
  "is_variant": false,          // 是否为变体
  "parent_id": "string"         // 母题ID
}
    ''')

    # 8. 测试需求
    doc.add_page_break()
    doc.add_heading('8. 测试需求', level=1)

    doc.add_heading('8.1 测试类型', level=2)
    test_table = doc.add_table(rows=5, cols=3)
    test_table.style = 'Table Grid'
    headers = ['类型', '覆盖率目标', '状态']
    for i, header in enumerate(headers):
        test_table.rows[0].cells[i].text = header
        test_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    tests = [
        ('单元测试', '≥ 80%', '部分完成'),
        ('集成测试', '核心流程', '已完成'),
        ('API测试', '所有端点', '待完善'),
        ('性能测试', '关键接口', '待实现')
    ]
    for i, (t_type, target, status) in enumerate(tests, 1):
        test_table.rows[i].cells[0].text = t_type
        test_table.rows[i].cells[1].text = target
        test_table.rows[i].cells[2].text = status

    # 9. 部署需求
    doc.add_page_break()
    doc.add_heading('9. 部署需求', level=1)

    doc.add_heading('9.1 环境要求', level=2)
    env_table = doc.add_table(rows=5, cols=3)
    env_table.style = 'Table Grid'
    headers = ['环境', '配置', '说明']
    for i, header in enumerate(headers):
        env_table.rows[0].cells[i].text = header
        env_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    envs = [
        ('Python', '3.10+', '运行环境'),
        ('SQLite', '-', '内置数据库'),
        ('内存', '≥ 2GB', '运行内存'),
        ('磁盘', '≥ 10GB', '数据存储')
    ]
    for i, (env, config, desc) in enumerate(envs, 1):
        env_table.rows[i].cells[0].text = env
        env_table.rows[i].cells[1].text = config
        env_table.rows[i].cells[2].text = desc

    doc.add_heading('9.2 启动方式', level=2)
    doc.add_paragraph('后端启动:', style='List Bullet')
    doc.add_paragraph('pip install -r requirements.txt', style='List Bullet 2')
    doc.add_paragraph('python launch.py', style='List Bullet 2')
    doc.add_paragraph('前端启动:', style='List Bullet')
    doc.add_paragraph('npm install', style='List Bullet 2')
    doc.add_paragraph('npm run dev', style='List Bullet 2')

    # 10. 待办事项
    doc.add_page_break()
    doc.add_heading('10. 待办事项 (Backlog)', level=1)

    doc.add_heading('10.1 高优先级', level=2)
    high_priority = [
        '实现 API 认证机制 (JWT)',
        '完善前端管理界面',
        '添加数据导出功能',
        '实现评测队列（异步处理）'
    ]
    for item in high_priority:
        doc.add_paragraph(f'[ ] {item}', style='List Bullet')

    doc.add_heading('10.2 中优先级', level=2)
    mid_priority = [
        '支持更多大模型平台（Claude、Gemini）',
        '添加评测报告可视化图表',
        '实现实时评测进度推送 (WebSocket)',
        '添加数据备份恢复功能'
    ]
    for item in mid_priority:
        doc.add_paragraph(f'[ ] {item}', style='List Bullet')

    doc.add_heading('10.3 低优先级', level=2)
    low_priority = [
        '支持自定义评分规则',
        '插件化评测维度',
        '多语言支持',
        '移动端适配'
    ]
    for item in low_priority:
        doc.add_paragraph(f'[ ] {item}', style='List Bullet')

    # 11. 附录
    doc.add_page_break()
    doc.add_heading('11. 附录', level=1)

    doc.add_heading('11.1 术语表', level=2)
    term_table = doc.add_table(rows=6, cols=2)
    term_table.style = 'Table Grid'
    headers = ['术语', '说明']
    for i, header in enumerate(headers):
        term_table.rows[0].cells[i].text = header
        term_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    terms = [
        ('Agent', '保险智能体，基于大语言模型的对话系统'),
        ('评测维度', '评估Agent能力的五个方向'),
        ('题目变异', '生成题目变体防止数据污染'),
        ('沙箱', '隔离的Agent执行环境'),
        ('Ground Truth', '标准答案')
    ]
    for i, (term, desc) in enumerate(terms, 1):
        term_table.rows[i].cells[0].text = term
        term_table.rows[i].cells[1].text = desc

    doc.add_heading('11.2 参考资料', level=2)
    doc.add_paragraph('• FastAPI 文档: https://fastapi.tiangolo.com', style='List Bullet')
    doc.add_paragraph('• Next.js 文档: https://nextjs.org/docs', style='List Bullet')

    # 保存文档
    output_path = '/Users/arthur/Apps/Projects/insurance-agent-benchmark/docs/PRD.docx'
    doc.save(output_path)
    print(f"✅ Word 文档已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        from docx import Document
        from docx.oxml import parse_xml
        create_prd_word()
    except ImportError:
        print("请先安装 python-docx:")
        print("pip install python-docx")
        print("\n或者使用以下命令生成 PDF:")
        print("pip install markdown weasyprint")
        print("python -m markdown docs/PRD.md -f pdf > docs/PRD.pdf")
